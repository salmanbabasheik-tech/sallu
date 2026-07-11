from http.server import BaseHTTPRequestHandler
import json
import boto3
import os
import base64
import re
from io import BytesIO

MODEL_ID       = "global.anthropic.claude-opus-4-6-v1"
AWS_ACCESS_KEY = os.environ.get('AWS_ACCESS_KEY_ID', '')
AWS_SECRET_KEY = os.environ.get('AWS_SECRET_ACCESS_KEY', '')
AWS_REGION     = os.environ.get('AWS_REGION', 'us-east-2')


def generate_text(prompt):
    client = boto3.client(
        'bedrock-runtime',
        region_name=AWS_REGION,
        aws_access_key_id=AWS_ACCESS_KEY,
        aws_secret_access_key=AWS_SECRET_KEY
    )
    body = json.dumps({
        "anthropic_version": "bedrock-2023-05-31",
        "max_tokens": 8096,
        "messages": [{"role": "user", "content": prompt}]
    })
    response = client.invoke_model(
        modelId=MODEL_ID,
        contentType="application/json",
        accept="application/json",
        body=body
    )
    result = json.loads(response['body'].read())
    return result['content'][0]['text']


def make_docx(title, content):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line in content.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\. ', line):
            doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
        else:
            p = doc.add_paragraph()
            parts = re.split(r'\*\*(.*?)\*\*', line)
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_pdf(title, content):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=2.5*cm, leftMargin=2.5*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm
    )
    styles = getSampleStyleSheet()
    story  = []

    story.append(Paragraph(title, styles['Title']))
    story.append(Spacer(1, 0.6*cm))

    for line in content.split('\n'):
        line = line.strip()
        if not line:
            story.append(Spacer(1, 0.25*cm))
            continue
        if line.startswith('# '):
            story.append(Paragraph(line[2:], styles['Heading1']))
        elif line.startswith('## '):
            story.append(Paragraph(line[3:], styles['Heading2']))
        elif line.startswith('### '):
            story.append(Paragraph(line[4:], styles['Heading3']))
        elif line.startswith('- ') or line.startswith('* '):
            story.append(Paragraph('• ' + line[2:], styles['Normal']))
        else:
            formatted = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
            story.append(Paragraph(formatted, styles['Normal']))
        story.append(Spacer(1, 0.15*cm))

    doc.build(story)
    return buf.getvalue()


class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        length = int(self.headers['Content-Length'])
        body   = json.loads(self.rfile.read(length).decode('utf-8'))

        prompt   = body.get('prompt', '')
        filetype = body.get('filetype', 'pdf').lower()
        title    = body.get('title', 'Generated Document')

        if not AWS_ACCESS_KEY or not AWS_SECRET_KEY:
            self._err('AWS credentials are not configured on the server.')
            return

        if not prompt:
            self._err('Prompt is required.')
            return

        try:
            content = generate_text(prompt)

            if filetype == 'docx':
                data = make_docx(title, content)
                mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                ext  = 'docx'
            else:
                data = make_pdf(title, content)
                mime = 'application/pdf'
                ext  = 'pdf'

            self.send_response(200)
            self.send_header('Content-Type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            self.wfile.write(json.dumps({
                'data': base64.b64encode(data).decode('utf-8'),
                'ext':  ext,
                'mime': mime,
                'title': title
            }).encode())

        except Exception as e:
            self._err(str(e))

    def _err(self, msg):
        self.send_response(500)
        self.send_header('Content-Type', 'application/json')
        self.send_header('Access-Control-Allow-Origin', '*')
        self.end_headers()
        self.wfile.write(json.dumps({'error': msg}).encode())

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()
