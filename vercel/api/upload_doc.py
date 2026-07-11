from http.server import BaseHTTPRequestHandler
import json
import boto3
import os
import base64

MODEL_ID       = "global.anthropic.claude-opus-4-6-v1"
AWS_ACCESS_KEY = os.environ.get('AWS_ACCESS_KEY_ID', '')
AWS_SECRET_KEY = os.environ.get('AWS_SECRET_ACCESS_KEY', '')
AWS_REGION     = os.environ.get('AWS_REGION', 'us-east-2')

PDF_TYPE  = 'application/pdf'
DOCX_TYPE = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
DOC_TYPE  = 'application/msword'
TEXT_TYPES = {'text/plain', 'text/html', 'text/markdown', 'text/csv'}
ALL_SUPPORTED = {PDF_TYPE, DOCX_TYPE, DOC_TYPE} | TEXT_TYPES


def extract_text_from_docx(raw_bytes):
    """Extract ALL text from docx including tables, headers, footers."""
    from docx import Document
    from io import BytesIO

    doc = Document(BytesIO(raw_bytes))
    lines = []

    # Extract header
    for section in doc.sections:
        for para in section.header.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())

    # Extract body — paragraphs AND tables interleaved in order
    for block in doc.element.body:
        tag = block.tag.split('}')[-1]

        if tag == 'p':
            # Regular paragraph
            from docx.oxml.ns import qn
            text = ''.join(node.text or '' for node in block.iter() if node.tag == qn('w:t'))
            if text.strip():
                lines.append(text.strip())

        elif tag == 'tbl':
            # Table — extract every cell
            from docx.oxml.ns import qn
            for row in block.iter(qn('w:tr')):
                row_cells = []
                for cell in row.iter(qn('w:tc')):
                    cell_text = ''.join(
                        node.text or '' for node in cell.iter() if node.tag == qn('w:t')
                    ).strip()
                    if cell_text:
                        row_cells.append(cell_text)
                if row_cells:
                    lines.append(' | '.join(row_cells))

    # Extract footer
    for section in doc.sections:
        for para in section.footer.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())

    return '\n'.join(lines)


def build_content_blocks(file_data_b64, file_type, file_name, question):
    raw_bytes = base64.b64decode(file_data_b64)

    # Strict instruction to prevent hallucination
    strict_note = (
        "\n\nCRITICAL INSTRUCTION: You must ONLY use information explicitly present in the "
        "document above. Do NOT invent, assume, or fill in any details that are not clearly "
        "stated in the document. If something is not in the document, say so explicitly. "
        "Never guess employer names, dates, years of experience, skills, or any other facts."
    )

    if file_type == PDF_TYPE:
        return [
            {
                "type": "document",
                "source": {
                    "type": "base64",
                    "media_type": "application/pdf",
                    "data": file_data_b64
                }
            },
            {"type": "text", "text": question + strict_note}
        ]

    elif file_type in (DOCX_TYPE, DOC_TYPE):
        try:
            text = extract_text_from_docx(raw_bytes)
            if not text.strip():
                raise ValueError("No text extracted")
        except Exception as e:
            text = f"[Extraction error: {e}]"

        wrapped = (
            f"=== START OF DOCUMENT: {file_name} ===\n\n"
            f"{text}\n\n"
            f"=== END OF DOCUMENT ===\n\n"
        )
        return [{"type": "text", "text": wrapped + question + strict_note}]

    else:
        try:
            text = raw_bytes.decode('utf-8', errors='replace')
        except Exception:
            text = "[Could not decode file]"
        wrapped = (
            f"=== START OF DOCUMENT: {file_name} ===\n\n"
            f"{text}\n\n"
            f"=== END OF DOCUMENT ===\n\n"
        )
        return [{"type": "text", "text": wrapped + question + strict_note}]


class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        content_length = int(self.headers['Content-Length'])
        post_data = self.rfile.read(content_length)
        body = json.loads(post_data.decode('utf-8'))

        file_data        = body.get('file_data', '')
        file_type        = body.get('file_type', '')
        file_name        = body.get('file_name', 'document')
        question         = body.get('question', 'Please analyze this document accurately.')
        # Do NOT pass old messages when analyzing a document — prevents context bleed
        messages_history = []

        if not AWS_ACCESS_KEY or not AWS_SECRET_KEY:
            self._stream_error('AWS credentials are not configured.')
            return
        if not file_data:
            self._stream_error('File data is required.')
            return
        if file_type not in ALL_SUPPORTED:
            self._stream_error(f'Unsupported file type: {file_type}.')
            return

        try:
            content_blocks = build_content_blocks(file_data, file_type, file_name, question)
        except Exception as e:
            self._stream_error(f'Failed to process file: {e}')
            return

        try:
            client = boto3.client(
                'bedrock-runtime',
                region_name=AWS_REGION,
                aws_access_key_id=AWS_ACCESS_KEY,
                aws_secret_access_key=AWS_SECRET_KEY
            )
            messages = [{"role": "user", "content": content_blocks}]

            request_body = json.dumps({
                "anthropic_version": "bedrock-2023-05-31",
                "max_tokens": 8096,
                "system": (
                    "You are XsAllu Ai, a precise document analysis assistant. "
                    "Your most important rule: NEVER invent, guess, or assume ANY information. "
                    "Only state facts that are explicitly written in the document provided. "
                    "If asked about something not in the document, say clearly: "
                    "'This information is not present in the document.' "
                    "Never hallucinate names, dates, companies, numbers, or any other details."
                ),
                "messages": messages
            })

            response = client.invoke_model_with_response_stream(
                modelId=MODEL_ID,
                contentType="application/json",
                accept="application/json",
                body=request_body
            )
        except Exception as e:
            self._stream_error(str(e))
            return

        self.send_response(200)
        self.send_header('Content-Type', 'application/x-ndjson')
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Cache-Control', 'no-cache')
        self.send_header('X-Accel-Buffering', 'no')
        self.end_headers()

        input_tokens = 0
        output_tokens = 0

        try:
            for event in response['body']:
                chunk = event.get('chunk')
                if not chunk:
                    continue
                payload = json.loads(chunk['bytes'])
                etype = payload.get('type')
                if etype == 'message_start':
                    input_tokens = payload.get('message', {}).get('usage', {}).get('input_tokens', 0)
                elif etype == 'content_block_delta':
                    delta = payload.get('delta', {})
                    if delta.get('type') == 'text_delta':
                        text_piece = delta.get('text', '')
                        if text_piece:
                            self.wfile.write((json.dumps({'type': 'chunk', 'text': text_piece}) + '\n').encode('utf-8'))
                            self.wfile.flush()
                elif etype == 'message_delta':
                    output_tokens = payload.get('usage', {}).get('output_tokens', output_tokens)

            self.wfile.write((json.dumps({
                'type': 'done',
                'usage': {'input_tokens': input_tokens, 'output_tokens': output_tokens}
            }) + '\n').encode('utf-8'))
            self.wfile.flush()
        except Exception as e:
            try:
                self.wfile.write((json.dumps({'type': 'error', 'error': str(e)}) + '\n').encode('utf-8'))
                self.wfile.flush()
            except Exception:
                pass

    def _stream_error(self, msg):
        self.send_response(500)
        self.send_header('Content-Type', 'application/x-ndjson')
        self.send_header('Access-Control-Allow-Origin', '*')
        self.end_headers()
        self.wfile.write((json.dumps({'type': 'error', 'error': msg}) + '\n').encode())

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()
